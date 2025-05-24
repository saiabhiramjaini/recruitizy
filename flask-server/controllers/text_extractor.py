import requests
import io
from pypdf import PdfReader
from docx import Document

# File processing functions
def extract_text_from_file(uploaded_file):
    file_extension = uploaded_file.filename.split('.')[-1].lower()
    if file_extension == "pdf":
        return extract_text_from_pdf(uploaded_file)
    elif file_extension == "docx":
        return extract_text_from_docx(uploaded_file)
    elif file_extension == "txt":
        return extract_text_from_txt(uploaded_file)
    return ""

def extract_text_from_pdf(pdf_file):
    text = ""
    pdf_reader = PdfReader(pdf_file)
    for page in pdf_reader.pages:
        text += page.extract_text() or ''
    return text

def extract_text_from_docx(docx_file):
    text = ""
    doc = Document(docx_file)
    for paragraph in doc.paragraphs:
        text += paragraph.text + '\n'
    return text

def extract_text_from_txt(txt_file):
    return txt_file.read().decode('utf-8')


def extract_text_from_s3_url(file_url: str) -> str:
    response = requests.get(file_url)
    if response.status_code != 200:
        raise Exception("Failed to download file from S3")
    
    content_type = response.headers.get("Content-Type", "")
    
    if "pdf" in content_type:
        # PDF file
        pdf_file = io.BytesIO(response.content)
        reader = PdfReader(pdf_file)
        return "\n".join([page.extract_text() for page in reader.pages])
    
    elif "word" in content_type or file_url.endswith(".docx"):
        import docx
        from docx import Document
        doc = Document(io.BytesIO(response.content))
        return "\n".join([para.text for para in doc.paragraphs])
    
    else:
        raise Exception("Unsupported file format")